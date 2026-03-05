#!/usr/bin/env python3
"""
SEO & SEM Audit Tool
====================
Crawls a URL, analyzes it with AI (Claude or Gemini), and generates a professional .docx report.

Usage:
    # Using Claude (default)
    python seo_audit.py <URL> --api-key sk-ant-...

    # Using Gemini
    python seo_audit.py <URL> --provider gemini --api-key AIzaSy...

    # Auto-detect provider from API key format
    python seo_audit.py <URL> --api-key <your-key>

Requirements:
    pip install requests beautifulsoup4 python-docx anthropic google-genai lxml
"""

import argparse
import json
import os
import re
import sys
import textwrap
from datetime import datetime
from urllib.parse import urlparse, urljoin

# Load .env file if present (requires: pip install python-dotenv)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv not installed — rely on shell environment variables

# ── Config from environment ────────────────────────────────────────
CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-20250514")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-pro-preview-03-25")

import requests
from bs4 import BeautifulSoup, Comment

# ═══════════════════════════════════════════════════════════════════
# FIXED SCORING RUBRIC — Never changes between runs.
# AI evaluates each check; Python computes scores from weights.
# PASS=1.0, WARN=0.5, FAIL=0.0
# ═══════════════════════════════════════════════════════════════════

TECHNICAL_CHECKS = [
    # (check_name, weight)
    # Higher weight = bigger impact on Technical score
    ("HTTPS / SSL",            8),   # Core trust signal
    ("Indexability",           8),   # meta robots, noindex detection
    ("Page Speed",             7),   # response_time_ms
    ("Mobile-Friendly",        7),   # has_viewport, responsive signals
    ("Canonical Tag",          6),   # duplicate URL prevention
    ("Structured Data",        6),   # JSON-LD schema types
    ("Mixed Content",          6),   # HTTP resources on HTTPS page
    ("Image Optimization",     6),   # WebP format, missing dimensions (CLS risk)
    ("HTTP Security Headers",  5),   # X-Frame-Options, CSP, HSTS, X-Content-Type
    ("Internal Linking",       5),   # internal link count and structure
    ("Duplicate Content",      5),   # repeated paragraphs detected
    ("Lazy Loading",           4),   # loading="lazy" on images
    ("Redirect Handling",      4),   # HTTP→HTTPS, www→non-www redirect chain
    ("Robots.txt",             4),   # exists, disallow rules, sitemap reference
    ("XML Sitemap",            4),   # exists, target URL included
    ("URL Structure",          3),   # clean slugs, no params/session IDs
]

CONTENT_CHECKS = [
    # (check_name, weight)
    ("Title Tag",              10),  # length 50–60 chars, keyword present
    ("H1 Tag",                 10),  # single H1, keyword-rich
    ("Meta Description",        8),  # length 120–160 chars, compelling
    ("Keyword Targeting",       8),  # target keyword in key positions
    ("Content Depth",           7),  # word count, paragraph quality
    ("Image Alt Text",          7),  # all images have descriptive alt text
    ("Heading Hierarchy",       6),  # H1→H2→H3 logical structure
    ("Keyword in URL",          5),  # target keyword present in URL slug
    ("OG / Social Tags",        5),  # og:title, og:description, og:image
    ("Local SEO / NAP",         5),  # Name, Address, Phone consistent
    ("External Links",          4),  # quality outbound links to authoritative sources
    ("FAQ / Rich Content",      4),  # structured FAQ, lists, tables
    ("CTA Placement",           4),  # clear calls-to-action present
    ("HTML Lang Attribute",     3),  # <html lang="en-au"> for local SEO + accessibility
]

SEM_CHECKS = [
    # (check_name, weight)
    ("Landing Page Relevance",   12),  # page matches ad intent
    ("Clear Value Proposition",  10),  # USP immediately visible
    ("Call to Action",           10),  # prominent, specific CTA button/text
    ("Above the Fold CTA",        8),  # CTA visible without scrolling
    ("Phone / Contact Visibility", 8), # phone number prominent (click-to-call)
    ("Lead Capture / Form",       7),  # contact form, booking, subscription
    ("Trust Signals",             7),  # accreditations, logos, certifications
    ("Social Proof / Reviews",    7),  # testimonials, ratings, case studies
    ("Ad Keyword Alignment",      7),  # landing page copy matches search terms
    ("Page Load Speed",           6),  # fast load = lower CPC, higher Quality Score
    ("Mobile UX",                 6),  # mobile layout, tap targets, readability
]

# Status score values
STATUS_VALUES = {"PASS": 1.0, "WARN": 0.5, "FAIL": 0.0}


def compute_scores(analysis: dict) -> dict:
    """
    Compute deterministic scores from fixed rubric weights.
    Overrides any AI-assigned scores so results are always comparable.
    """
    def _score_section(checks_rubric, checks_data):
        lookup = {c.get("name", "").strip(): c for c in checks_data}
        total_weight = sum(w for _, w in checks_rubric)
        earned = 0.0
        for name, weight in checks_rubric:
            check = lookup.get(name, {})
            status = check.get("status", "FAIL").upper()
            earned += weight * STATUS_VALUES.get(status, 0.0)
        return round(earned / total_weight * 100) if total_weight else 0

    tech_score = _score_section(TECHNICAL_CHECKS, analysis.get("technical_seo", {}).get("checks", []))
    content_score = _score_section(CONTENT_CHECKS, analysis.get("content_seo", {}).get("checks", []))
    sem_score = _score_section(SEM_CHECKS, analysis.get("sem_readiness", {}).get("checks", []))

    # Overall = weighted average (tech 40%, content 35%, sem 25%)
    overall = round(tech_score * 0.40 + content_score * 0.35 + sem_score * 0.25)

    def _grade(score):
        if score >= 90: return "A"
        if score >= 80: return "A-" if score >= 85 else "B+"
        if score >= 70: return "B"
        if score >= 60: return "C"
        if score >= 50: return "D"
        return "F"

    analysis["technical_seo"]["score"] = tech_score
    analysis["technical_seo"]["grade"] = _grade(tech_score)
    analysis["content_seo"]["score"] = content_score
    analysis["content_seo"]["grade"] = _grade(content_score)
    analysis["sem_readiness"]["score"] = sem_score
    analysis["sem_readiness"]["grade"] = _grade(sem_score)
    analysis["overall_score"] = overall
    analysis["overall_grade"] = _grade(overall)
    return analysis


def _history_path(url: str) -> str:
    """Return the history JSON file path for a given domain."""
    domain = urlparse(url).netloc.replace("www.", "").replace(".", "-")
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), f"seo-history-{domain}.json")


def load_history(url: str) -> list:
    """Load previous audit runs for this domain."""
    path = _history_path(url)
    if os.path.exists(path):
        try:
            with open(path) as f:
                return json.load(f)
        except Exception:
            return []
    return []


def save_history(url: str, analysis: dict):
    """Append current run results to history file."""
    history = load_history(url)
    entry = {
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "url": url,
        "overall_score": analysis.get("overall_score", 0),
        "overall_grade": analysis.get("overall_grade", "N/A"),
        "technical_score": analysis.get("technical_seo", {}).get("score", 0),
        "technical_grade": analysis.get("technical_seo", {}).get("grade", "N/A"),
        "content_score": analysis.get("content_seo", {}).get("score", 0),
        "content_grade": analysis.get("content_seo", {}).get("grade", "N/A"),
        "sem_score": analysis.get("sem_readiness", {}).get("score", 0),
        "sem_grade": analysis.get("sem_readiness", {}).get("grade", "N/A"),
        "check_results": {
            "technical": {c.get("name"): c.get("status") for c in analysis.get("technical_seo", {}).get("checks", [])},
            "content": {c.get("name"): c.get("status") for c in analysis.get("content_seo", {}).get("checks", [])},
            "sem": {c.get("name"): c.get("status") for c in analysis.get("sem_readiness", {}).get("checks", [])},
        },
    }
    history.append(entry)
    with open(_history_path(url), "w") as f:
        json.dump(history, f, indent=2)
    print(f"   📊 History saved: {_history_path(url)}")


# ═══════════════════════════════════════════════════════════════════
# PART 1: WEB CRAWLER — Extracts all SEO-relevant data from a URL
# ═══════════════════════════════════════════════════════════════════

class SEOCrawler:
    """Crawls a URL and extracts all SEO-relevant signals."""

    HEADERS = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }

    def __init__(self, url: str):
        self.url = url
        self.parsed_url = urlparse(url)
        self.base_url = f"{self.parsed_url.scheme}://{self.parsed_url.netloc}"
        self.data = {}

    def crawl(self) -> dict:
        """Run full crawl and return structured data."""
        print(f"\n🔍 Crawling: {self.url}")
        print("=" * 60)

        self._fetch_page()
        self._check_robots_txt()
        self._check_sitemap()
        self._extract_meta()
        self._extract_headings()
        self._extract_content()
        self._extract_links()
        self._extract_images()
        self._extract_schema()
        self._extract_forms()
        self._extract_social_tags()
        self._check_https()
        self._check_canonical()
        self._extract_scripts_and_styles()

        print(f"\n✅ Crawl complete — {len(self.data)} data points collected")
        return self.data

    def _fetch_page(self):
        """Fetch the page HTML."""
        print("  📄 Fetching page...")
        try:
            resp = requests.get(self.url, headers=self.HEADERS, timeout=30, allow_redirects=True)
            self.data["status_code"] = resp.status_code
            self.data["final_url"] = resp.url
            self.data["response_time_ms"] = round(resp.elapsed.total_seconds() * 1000)
            self.data["content_length"] = len(resp.content)
            self.data["headers"] = dict(resp.headers)
            self.html = resp.text
            self.soup = BeautifulSoup(self.html, "lxml")
            print(f"    Status: {resp.status_code} | Size: {len(resp.content):,} bytes | Time: {self.data['response_time_ms']}ms")
        except Exception as e:
            print(f"    ❌ Failed to fetch: {e}")
            self.data["fetch_error"] = str(e)
            self.html = ""
            self.soup = BeautifulSoup("", "lxml")

    def _check_robots_txt(self):
        """Check robots.txt."""
        print("  🤖 Checking robots.txt...")
        try:
            resp = requests.get(f"{self.base_url}/robots.txt", headers=self.HEADERS, timeout=10)
            self.data["robots_txt_status"] = resp.status_code
            if resp.status_code == 200:
                self.data["robots_txt_content"] = resp.text[:2000]
                self.data["robots_txt_has_sitemap"] = "sitemap" in resp.text.lower()
                print(f"    Found ({len(resp.text)} chars)")
            else:
                print(f"    Status: {resp.status_code}")
        except Exception as e:
            self.data["robots_txt_status"] = "error"
            print(f"    ❌ Error: {e}")

    def _check_sitemap(self):
        """Check sitemap.xml."""
        print("  🗺️  Checking sitemap...")
        try:
            resp = requests.get(f"{self.base_url}/sitemap.xml", headers=self.HEADERS, timeout=10)
            self.data["sitemap_status"] = resp.status_code
            if resp.status_code == 200:
                sitemap_soup = BeautifulSoup(resp.text, features="xml")
                urls = sitemap_soup.find_all("loc")
                self.data["sitemap_url_count"] = len(urls)
                self.data["sitemap_contains_target"] = any(self.url.rstrip("/") in u.text for u in urls)
                print(f"    Found ({len(urls)} URLs, target page {'included' if self.data['sitemap_contains_target'] else 'NOT included'})")
            else:
                print(f"    Status: {resp.status_code}")
        except Exception as e:
            self.data["sitemap_status"] = "error"
            print(f"    ❌ Error: {e}")

    def _extract_meta(self):
        """Extract meta tags."""
        print("  🏷️  Extracting meta tags...")
        # Title
        title_tag = self.soup.find("title")
        self.data["title"] = title_tag.text.strip() if title_tag else None
        self.data["title_length"] = len(self.data["title"]) if self.data["title"] else 0

        # Meta description
        meta_desc = self.soup.find("meta", attrs={"name": "description"})
        self.data["meta_description"] = meta_desc["content"].strip() if meta_desc and meta_desc.get("content") else None
        self.data["meta_description_length"] = len(self.data["meta_description"]) if self.data["meta_description"] else 0

        # Meta robots
        meta_robots = self.soup.find("meta", attrs={"name": "robots"})
        self.data["meta_robots"] = meta_robots["content"] if meta_robots and meta_robots.get("content") else None

        # Viewport
        viewport = self.soup.find("meta", attrs={"name": "viewport"})
        self.data["has_viewport"] = viewport is not None

        # Charset
        charset = self.soup.find("meta", attrs={"charset": True})
        self.data["charset"] = charset["charset"] if charset else None

        # HTML lang attribute
        html_tag = self.soup.find("html")
        self.data["html_lang"] = html_tag.get("lang") if html_tag else None

        print(f"    Title: '{self.data['title']}' ({self.data['title_length']} chars)")
        print(f"    Meta desc: {'Set' if self.data['meta_description'] else 'NOT SET'} ({self.data['meta_description_length']} chars)")
        print(f"    HTML lang: {self.data['html_lang'] or 'NOT SET'}")

    def _extract_headings(self):
        """Extract heading hierarchy."""
        print("  📝 Extracting headings...")
        headings = {}
        for level in range(1, 7):
            tags = self.soup.find_all(f"h{level}")
            headings[f"h{level}"] = [tag.get_text(strip=True) for tag in tags]
        self.data["headings"] = headings
        for level, texts in headings.items():
            if texts:
                print(f"    {level.upper()}: {len(texts)} found — {texts[0][:60]}{'...' if len(texts[0]) > 60 else ''}")

    def _extract_content(self):
        """Extract and analyze body content."""
        print("  📖 Analyzing content...")
        # Get text content
        for elem in self.soup(["script", "style", "nav", "header", "footer"]):
            elem.decompose()
        body = self.soup.find("body")
        text = body.get_text(separator=" ", strip=True) if body else ""
        words = text.split()
        self.data["word_count"] = len(words)
        self.data["content_text"] = text[:5000]  # First 5000 chars for analysis

        # Check for duplicate paragraphs
        paragraphs = [p.get_text(strip=True) for p in self.soup.find_all("p") if len(p.get_text(strip=True)) > 50]
        seen = {}
        duplicates = []
        for p in paragraphs:
            if p in seen:
                duplicates.append(p[:100] + "...")
            seen[p] = True
        self.data["duplicate_paragraphs"] = duplicates
        self.data["paragraph_count"] = len(paragraphs)

        print(f"    Word count: {self.data['word_count']}")
        print(f"    Paragraphs: {self.data['paragraph_count']}")
        if duplicates:
            print(f"    ⚠️  Duplicate paragraphs found: {len(duplicates)}")

    def _extract_links(self):
        """Extract and categorize links."""
        # Re-parse since we decomposed elements above
        soup = BeautifulSoup(self.html, "lxml")
        print("  🔗 Extracting links...")
        internal = []
        external = []
        tel_links = []
        broken_hrefs = []
        for a in soup.find_all("a", href=True):
            href = a["href"]
            text = a.get_text(strip=True)
            if href.startswith("tel:"):
                tel_links.append(href.replace("tel:", "").strip())
                continue
            if href.startswith("#") or href.startswith("mailto:"):
                continue
            full_url = urljoin(self.url, href)
            parsed = urlparse(full_url)
            entry = {"href": full_url, "text": text[:80], "original_href": href}
            if parsed.netloc == self.parsed_url.netloc or parsed.netloc == "":
                internal.append(entry)
            else:
                external.append(entry)

        self.data["internal_links"] = internal[:50]  # Cap for API
        self.data["internal_link_count"] = len(internal)
        self.data["external_links"] = external[:20]
        self.data["external_link_count"] = len(external)
        self.data["tel_links"] = tel_links  # Phone numbers found in tel: hrefs
        self.data["has_phone_link"] = len(tel_links) > 0
        print(f"    Internal: {len(internal)} | External: {len(external)} | Phone links: {len(tel_links)}")

    def _extract_images(self):
        """Extract image information."""
        soup = BeautifulSoup(self.html, "lxml")
        print("  🖼️  Extracting images...")
        images = []
        for img in soup.find_all("img"):
            src = img.get("src", "")
            alt = img.get("alt", "")
            width = img.get("width")
            height = img.get("height")
            loading = img.get("loading")
            images.append({
                "src": src[:200],
                "alt": alt[:150],
                "has_alt": bool(alt.strip()),
                "has_dimensions": bool(width and height),
                "has_lazy_loading": loading == "lazy",
                "format": src.split(".")[-1].split("?")[0].lower() if "." in src else "unknown",
            })
        self.data["images"] = images[:30]
        self.data["image_count"] = len(images)
        missing_alt = sum(1 for img in images if not img["has_alt"])
        png_count = sum(1 for img in images if img["format"] == "png")
        lazy_count = sum(1 for img in images if img["has_lazy_loading"])
        print(f"    Total: {len(images)} | Missing alt: {missing_alt} | PNGs: {png_count} | Lazy-loaded: {lazy_count}")

    def _extract_schema(self):
        """Extract structured data (JSON-LD)."""
        soup = BeautifulSoup(self.html, "lxml")
        print("  📊 Checking structured data...")
        schemas = []
        for script in soup.find_all("script", type="application/ld+json"):
            try:
                data = json.loads(script.string)
                if isinstance(data, list):
                    for item in data:
                        schemas.append(item.get("@type", "Unknown"))
                else:
                    schemas.append(data.get("@type", "Unknown"))
            except (json.JSONDecodeError, TypeError):
                pass
        self.data["schema_types"] = schemas
        self.data["has_schema"] = len(schemas) > 0
        print(f"    {'Found: ' + ', '.join(schemas) if schemas else 'No JSON-LD schema detected'}")

    def _extract_forms(self):
        """Extract form information."""
        soup = BeautifulSoup(self.html, "lxml")
        print("  📋 Checking forms...")
        forms = []
        for form in soup.find_all("form"):
            fields = []
            for inp in form.find_all(["input", "select", "textarea"]):
                name = inp.get("name", inp.get("id", "unnamed"))
                field_type = inp.get("type", inp.name)
                fields.append({"name": name, "type": field_type})
            forms.append({"action": form.get("action", ""), "method": form.get("method", ""), "field_count": len(fields), "fields": fields[:20]})
        self.data["forms"] = forms
        self.data["form_count"] = len(forms)
        print(f"    Forms found: {len(forms)}")

    def _extract_social_tags(self):
        """Extract Open Graph and Twitter card tags."""
        soup = BeautifulSoup(self.html, "lxml")
        print("  📱 Checking social/OG tags...")
        og_tags = {}
        for meta in soup.find_all("meta", attrs={"property": re.compile(r"^og:")}):
            og_tags[meta["property"]] = meta.get("content", "")[:200]
        twitter_tags = {}
        for meta in soup.find_all("meta", attrs={"name": re.compile(r"^twitter:")}):
            twitter_tags[meta["name"]] = meta.get("content", "")[:200]
        self.data["og_tags"] = og_tags
        self.data["twitter_tags"] = twitter_tags
        self.data["has_og_tags"] = len(og_tags) > 0
        self.data["has_twitter_tags"] = len(twitter_tags) > 0
        print(f"    OG tags: {'Yes (' + str(len(og_tags)) + ')' if og_tags else 'None'} | Twitter: {'Yes' if twitter_tags else 'None'}")

    def _check_https(self):
        """Check HTTPS and SSL."""
        print("  🔒 Checking HTTPS...")
        self.data["is_https"] = self.parsed_url.scheme == "https"
        # Check for mixed content (images/scripts via src, stylesheets via href)
        soup = BeautifulSoup(self.html, "lxml")
        mixed = []
        for tag in soup.find_all(["img", "script"]):
            src = tag.get("src", "")
            if src.startswith("http://"):
                mixed.append(src[:100])
        for tag in soup.find_all("link", rel=lambda r: r and "stylesheet" in r):
            href = tag.get("href", "")
            if href.startswith("http://"):
                mixed.append(href[:100])
        self.data["mixed_content"] = mixed[:10]
        self.data["has_mixed_content"] = len(mixed) > 0
        print(f"    HTTPS: {'Yes' if self.data['is_https'] else 'No'} | Mixed content: {len(mixed)} items")

    def _check_canonical(self):
        """Check canonical tag."""
        soup = BeautifulSoup(self.html, "lxml")
        print("  🔖 Checking canonical...")
        canonical = soup.find("link", rel="canonical")
        self.data["canonical_url"] = canonical["href"] if canonical and canonical.get("href") else None
        self.data["has_canonical"] = canonical is not None
        print(f"    Canonical: {self.data['canonical_url'] or 'Not set'}")

    def _extract_scripts_and_styles(self):
        """Count external scripts and stylesheets."""
        soup = BeautifulSoup(self.html, "lxml")
        scripts = soup.find_all("script", src=True)
        styles = soup.find_all("link", rel="stylesheet")
        self.data["external_script_count"] = len(scripts)
        self.data["external_style_count"] = len(styles)

        # Detect CMS / tech
        indicators = []
        if 'wp-content' in self.html:
            indicators.append("WordPress")
        if 'shopify' in self.html.lower():
            indicators.append("Shopify")
        if 'wix.com' in self.html.lower():
            indicators.append("Wix")
        if 'squarespace' in self.html.lower():
            indicators.append("Squarespace")
        if 'fbq(' in self.html:
            indicators.append("Facebook Pixel")
        if 'gtag(' in self.html or 'google-analytics' in self.html:
            indicators.append("Google Analytics")
        if 'googletagmanager' in self.html:
            indicators.append("Google Tag Manager")
        self.data["tech_detected"] = indicators
        print(f"  ⚙️  Tech detected: {', '.join(indicators) if indicators else 'None identified'}")


# ═══════════════════════════════════════════════════════════════════
# PART 2: AI ANALYZER — Sends crawl data to Claude or Gemini for analysis
# ═══════════════════════════════════════════════════════════════════

def _build_checklist_prompt() -> str:
    """Build the fixed checklist section of the prompt from the scoring rubric."""
    tech_names = "\n".join(f'      "{name}"' for name, _ in TECHNICAL_CHECKS)
    content_names = "\n".join(f'      "{name}"' for name, _ in CONTENT_CHECKS)
    sem_names = "\n".join(f'      "{name}"' for name, _ in SEM_CHECKS)
    return f"""
TECHNICAL SEO — evaluate EXACTLY these {len(TECHNICAL_CHECKS)} checks (exact names, no additions, no omissions):
{tech_names}

CONTENT SEO — evaluate EXACTLY these {len(CONTENT_CHECKS)} checks (exact names, no additions, no omissions):
{content_names}

SEM READINESS — evaluate EXACTLY these {len(SEM_CHECKS)} checks (exact names, no additions, no omissions):
{sem_names}
"""


ANALYSIS_PROMPT = """You are an expert SEO/SEM auditor. I have crawled a website and extracted the following technical data. Based on this data, produce a comprehensive SEO & SEM audit.

## CRAWL DATA
```json
{crawl_data}
```

## FIXED CHECKLIST
You MUST evaluate exactly the checks listed below. Do NOT add new checks, rename checks, or skip any check. The check names must match EXACTLY as written.

{checklist}

## YOUR TASK
Respond ONLY with a valid JSON object (no markdown, no backticks, no preamble):

{{
  "business_name": "detected business name",
  "business_description": "1-sentence description of what the business does",
  "executive_summary": "2-3 paragraphs covering key strengths and opportunities. Reference actual crawl data.",
  "key_strengths": "comma-separated list of strengths",
  "key_opportunities": "comma-separated list of opportunities",
  "technical_seo": {{
    "score": 0,
    "grade": "N/A",
    "checks": [
      {{
        "name": "EXACT check name from the fixed list above",
        "status": "PASS or WARN or FAIL",
        "finding": "Specific finding referencing actual crawl data",
        "recommendation": "Actionable recommendation"
      }}
    ],
    "priority_actions": ["action 1 - be specific and actionable", "action 2", "action 3", "action 4", "action 5"]
  }},
  "content_seo": {{
    "score": 0,
    "grade": "N/A",
    "checks": [
      {{
        "name": "EXACT check name from the fixed list above",
        "status": "PASS or WARN or FAIL",
        "finding": "Specific finding referencing actual crawl data",
        "recommendation": "Actionable recommendation"
      }}
    ],
    "priority_actions": ["action 1", "action 2", "action 3", "action 4", "action 5"]
  }},
  "sem_readiness": {{
    "score": 0,
    "grade": "N/A",
    "checks": [
      {{
        "name": "EXACT check name from the fixed list above",
        "status": "PASS or WARN or FAIL",
        "finding": "Specific finding referencing actual crawl data",
        "recommendation": "Actionable recommendation"
      }}
    ],
    "strengths": ["strength 1", "strength 2", "strength 3"],
    "issues": ["issue 1", "issue 2", "issue 3"],
    "ad_groups": [
      {{
        "name": "Ad Group Name",
        "keywords": "keyword 1, keyword 2, keyword 3, keyword 4",
        "rationale": "Why this page is the right destination for these keywords"
      }}
    ],
    "campaign_recommendations": ["recommendation 1", "recommendation 2", "recommendation 3", "recommendation 4", "recommendation 5"]
  }},
  "overall_score": 0,
  "overall_grade": "N/A",
  "quick_wins": [
    {{
      "rank": 1,
      "action": "What to do — be specific",
      "impact": "High/Medium/Low — with brief explanation",
      "effort": "High/Medium/Low — with brief explanation"
    }}
  ]
}}

CRITICAL RULES:
- The "checks" arrays must contain EXACTLY the named checks listed — same names, same count, no extras
- score and grade fields are placeholders (they will be calculated automatically) — set them to 0 and "N/A"
- Include AT LEAST 4 ad groups with 4+ keywords each
- Include AT LEAST 5 quick wins ranked by impact vs effort
- Reference ACTUAL data from the crawl (real title tag, real headings, real URLs, etc.)
- Be harsh but fair — do not inflate status; use FAIL if something is missing, WARN if partial/suboptimal
- Respond with ONLY the JSON object, nothing else
"""


def _prepare_prompt(crawl_data: dict) -> str:
    """Prepare the analysis prompt with crawl data and fixed checklist."""
    trimmed = {k: v for k, v in crawl_data.items()}
    if "content_text" in trimmed and len(str(trimmed["content_text"])) > 3000:
        trimmed["content_text"] = str(trimmed["content_text"])[:3000] + "...[truncated]"
    prompt = ANALYSIS_PROMPT.replace("{crawl_data}", json.dumps(trimmed, indent=2, default=str))
    prompt = prompt.replace("{checklist}", _build_checklist_prompt())
    return prompt


def _parse_json_response(text: str) -> dict:
    """Parse JSON from an AI response, handling markdown fences."""
    text = text.strip()
    text = re.sub(r"^```json\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("Could not find JSON in AI response. Raw response:\n" + text[:500])
    return json.loads(text[start:end + 1])


def analyze_with_claude(crawl_data: dict, api_key: str) -> dict:
    """Send crawl data to Claude API for expert analysis."""
    import anthropic

    print(f"\n🤖 Sending crawl data to Claude ({CLAUDE_MODEL}) for analysis...")
    print("   (This may take 30-60 seconds)")

    prompt = _prepare_prompt(crawl_data)
    client = anthropic.Anthropic(api_key=api_key)

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}],
    )

    text = ""
    for block in message.content:
        if hasattr(block, "text"):
            text += block.text

    try:
        result = _parse_json_response(text)
    except (json.JSONDecodeError, ValueError) as e:
        print(f"\n❌ Claude returned invalid JSON: {e}")
        print("   Raw response (first 500 chars):", text[:500])
        sys.exit(1)
    print("   ✅ Analysis complete!")
    return result


def analyze_with_gemini(crawl_data: dict, api_key: str) -> dict:
    """Send crawl data to Google Gemini API for expert analysis."""
    from google import genai

    print(f"\n🤖 Sending crawl data to Gemini ({GEMINI_MODEL}) for analysis...")
    print("   (This may take 30-60 seconds)")

    prompt = _prepare_prompt(crawl_data)
    client = genai.Client(api_key=api_key)

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
    )

    text = response.text
    try:
        result = _parse_json_response(text)
    except (json.JSONDecodeError, ValueError) as e:
        print(f"\n❌ Gemini returned invalid JSON: {e}")
        print("   Raw response (first 500 chars):", text[:500])
        sys.exit(1)
    print("   ✅ Analysis complete!")
    return result


def detect_provider(api_key: str) -> str:
    """Auto-detect AI provider from API key format."""
    if api_key.startswith("sk-ant-"):
        return "claude"
    elif api_key.startswith("AIzaSy"):
        return "gemini"
    else:
        # Could be either — default to claude
        return "claude"


# ═══════════════════════════════════════════════════════════════════
# PART 3: REPORT GENERATOR — Creates a formatted .docx report
# ═══════════════════════════════════════════════════════════════════

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Color constants
BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x1B, 0x1B, 0x1B)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_HEX = "27AE60"
AMBER_HEX = "F39C12"
RED_HEX = "E74C3C"
BLUE_HEX = "1F4E79"
LIGHT_BLUE_HEX = "D6E4F0"


def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def status_color_hex(status):
    if status == "PASS":
        return GREEN_HEX
    elif status == "WARN":
        return AMBER_HEX
    return RED_HEX


def add_styled_paragraph(doc, text, font_size=11, bold=False, color=GRAY, alignment=None, space_after=6):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    run.bold = bold
    run.font.color.rgb = color
    return p


def add_bold_body(doc, label, text):
    """Add paragraph with bold label + normal text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(label)
    r1.font.size = Pt(11)
    r1.font.name = "Arial"
    r1.bold = True
    r1.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = "Arial"
    r2.font.color.rgb = GRAY
    return p


def style_header_cell(cell, text):
    """Style a table header cell."""
    set_cell_shading(cell, BLUE_HEX)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.bold = True
    run.font.color.rgb = WHITE


def style_cell(cell, text, bold=False, color=DARK, alignment=None):
    """Style a regular table cell."""
    p = cell.paragraphs[0]
    p.clear()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.bold = bold
    run.font.color.rgb = color


def style_status_cell(cell, status):
    """Style a status cell with color background."""
    set_cell_shading(cell, status_color_hex(status))
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(status)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.bold = True
    run.font.color.rgb = WHITE


def generate_report(analysis: dict, url: str, output_path: str, previous_run: dict = None):
    """Generate a .docx report from analysis data."""
    print(f"\n📄 Generating report: {output_path}")

    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── TITLE PAGE ──
    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(doc, "SEO & SEM Strategy Report", 28, True, BLUE, WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, f"{analysis.get('business_name', 'Website')} — Homepage Audit", 14, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, url, 11, False, BLUE, WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, f"Report Date: {datetime.now().strftime('%d %B %Y')}", 11, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ──
    doc.add_heading("Executive Summary", level=1)
    for para in analysis.get("executive_summary", "").split("\n"):
        if para.strip():
            add_styled_paragraph(doc, para.strip(), 11, False, GRAY)

    add_bold_body(doc, "Key strengths: ", analysis.get("key_strengths", ""))
    add_bold_body(doc, "Key opportunities: ", analysis.get("key_opportunities", ""))

    # ── SCORES TABLE ──
    doc.add_heading("Overall Audit Scores", level=2)
    def _delta_str(current, prev_key):
        """Format score change vs previous run."""
        if not previous_run:
            return ""
        prev = previous_run.get(prev_key)
        if prev is None:
            return ""
        diff = current - prev
        if diff > 0:
            return f"  (+{diff})"
        elif diff < 0:
            return f"  ({diff})"
        return "  (=)"

    tech_score = analysis.get("technical_seo", {}).get("score", 0)
    content_score = analysis.get("content_seo", {}).get("score", 0)
    sem_score = analysis.get("sem_readiness", {}).get("score", 0)
    overall_score = analysis.get("overall_score", 0)

    scores_data = [
        ("Technical SEO",           tech_score,    analysis.get("technical_seo", {}).get("grade", "N/A"),    _delta_str(tech_score, "technical_score")),
        ("Content SEO",             content_score, analysis.get("content_seo", {}).get("grade", "N/A"),      _delta_str(content_score, "content_score")),
        ("SEM / AdWords Readiness", sem_score,     analysis.get("sem_readiness", {}).get("grade", "N/A"),    _delta_str(sem_score, "sem_score")),
    ]

    has_history = previous_run is not None
    col_count = 4 if has_history else 3
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    style_header_cell(hdr[0], "Category")
    style_header_cell(hdr[1], "Score")
    style_header_cell(hdr[2], "Grade")
    if has_history:
        style_header_cell(hdr[3], f"vs {previous_run.get('date', 'prev')[:10]}")

    for cat, score, grade, delta in scores_data:
        row = table.add_row().cells
        style_cell(row[0], cat)
        style_cell(row[1], f"{score} / 100", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell(row[2], grade, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if has_history:
            delta_color = GREEN_HEX if "+" in delta else (RED_HEX if "-" in delta else "888888")
            set_cell_shading(row[3], delta_color)
            style_cell(row[3], delta.strip(), alignment=WD_ALIGN_PARAGRAPH.CENTER, color=WHITE)

    # Overall row
    row = table.add_row().cells
    style_cell(row[0], "Overall", bold=True)
    set_cell_shading(row[0], LIGHT_BLUE_HEX)
    style_cell(row[1], f"{overall_score} / 100", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(row[1], LIGHT_BLUE_HEX)
    style_cell(row[2], analysis.get("overall_grade", "N/A"), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(row[2], LIGHT_BLUE_HEX)
    if has_history:
        overall_delta = _delta_str(overall_score, "overall_score")
        delta_color = GREEN_HEX if "+" in overall_delta else (RED_HEX if "-" in overall_delta else "888888")
        set_cell_shading(row[3], delta_color)
        style_cell(row[3], overall_delta.strip(), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=WHITE)

    doc.add_page_break()

    # ── 1. TECHNICAL SEO ──
    doc.add_heading("1. Technical SEO Review", level=1)
    add_styled_paragraph(doc, "Is the page indexed correctly and technically sound for organic search?")

    tech_checks = analysis.get("technical_seo", {}).get("checks", [])
    if tech_checks:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        style_header_cell(hdr[0], "Check")
        style_header_cell(hdr[1], "Status")
        style_header_cell(hdr[2], "Finding & Recommendation")

        for check in tech_checks:
            row = table.add_row().cells
            style_cell(row[0], check.get("name", ""), bold=True)
            style_status_cell(row[1], check.get("status", "WARN"))
            finding = check.get("finding", "")
            rec = check.get("recommendation", "")
            style_cell(row[2], f"{finding} RECOMMENDATION: {rec}" if rec else finding)

    doc.add_heading("Technical SEO — Priority Actions", level=2)
    for i, action in enumerate(analysis.get("technical_seo", {}).get("priority_actions", []), 1):
        add_styled_paragraph(doc, f"{i}. {action}")

    doc.add_page_break()

    # ── 2. CONTENT SEO ──
    doc.add_heading("2. Content SEO Review", level=1)
    add_styled_paragraph(doc, "Is the content appropriate for organic traffic to find this page?")

    content_checks = analysis.get("content_seo", {}).get("checks", [])
    if content_checks:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        style_header_cell(hdr[0], "Check")
        style_header_cell(hdr[1], "Status")
        style_header_cell(hdr[2], "Finding & Recommendation")

        for check in content_checks:
            row = table.add_row().cells
            style_cell(row[0], check.get("name", ""), bold=True)
            style_status_cell(row[1], check.get("status", "WARN"))
            finding = check.get("finding", "")
            rec = check.get("recommendation", "")
            style_cell(row[2], f"{finding} RECOMMENDATION: {rec}" if rec else finding)

    doc.add_heading("Content SEO — Priority Actions", level=2)
    for i, action in enumerate(analysis.get("content_seo", {}).get("priority_actions", []), 1):
        add_styled_paragraph(doc, f"{i}. {action}")

    doc.add_page_break()

    # ── 3. SEM / ADWORDS ──
    doc.add_heading("3. Google Ads / SEM — Destination URL Review", level=1)

    sem = analysis.get("sem_readiness", {})
    add_styled_paragraph(doc, f"Overall SEM Readiness Score: {sem.get('score', 0)}/100 — {sem.get('grade', 'N/A')}", 11, True)

    sem_checks = sem.get("checks", [])
    if sem_checks:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        style_header_cell(hdr[0], "Check")
        style_header_cell(hdr[1], "Status")
        style_header_cell(hdr[2], "Finding & Recommendation")
        for check in sem_checks:
            row = table.add_row().cells
            style_cell(row[0], check.get("name", ""), bold=True)
            style_status_cell(row[1], check.get("status", "WARN"))
            finding = check.get("finding", "")
            rec = check.get("recommendation", "")
            style_cell(row[2], f"{finding} RECOMMENDATION: {rec}" if rec else finding)
        doc.add_paragraph()

    add_styled_paragraph(doc, "Strengths as a paid landing page:", 11, True, DARK)
    for s in sem.get("strengths", []):
        add_styled_paragraph(doc, f"• {s}")

    add_styled_paragraph(doc, "Issues that reduce paid campaign performance:", 11, True, DARK)
    for s in sem.get("issues", []):
        add_styled_paragraph(doc, f"• {s}")

    # Ad Groups table
    doc.add_heading("Recommended Ad Groups — Destination: This Page", level=2)
    ad_groups = sem.get("ad_groups", [])
    if ad_groups:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        style_header_cell(hdr[0], "Ad Group")
        style_header_cell(hdr[1], "Target Keywords")
        style_header_cell(hdr[2], "Why This Page")

        for group in ad_groups:
            row = table.add_row().cells
            style_cell(row[0], group.get("name", ""), bold=True)
            style_cell(row[1], group.get("keywords", ""))
            style_cell(row[2], group.get("rationale", ""))

    doc.add_heading("SEM Campaign Recommendations", level=2)
    for i, rec in enumerate(sem.get("campaign_recommendations", []), 1):
        add_styled_paragraph(doc, f"{i}. {rec}")

    doc.add_page_break()

    # ── 4. QUICK WINS ──
    doc.add_heading("4. Top Quick Wins (Ranked by Impact vs Effort)", level=1)

    qw = analysis.get("quick_wins", [])
    if qw:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        style_header_cell(hdr[0], "#")
        style_header_cell(hdr[1], "Action")
        style_header_cell(hdr[2], "Impact")
        style_header_cell(hdr[3], "Effort")

        for win in qw:
            row = table.add_row().cells
            style_cell(row[0], str(win.get("rank", "")), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            style_cell(row[1], win.get("action", ""), bold=True)
            style_cell(row[2], win.get("impact", ""))
            style_cell(row[3], win.get("effort", ""))

    # Footer note
    doc.add_paragraph()
    add_styled_paragraph(doc, f"Report prepared based on live page crawl of {url} conducted on {datetime.now().strftime('%d %B %Y')}.", 10, False, GRAY)
    add_styled_paragraph(doc, "For questions or follow-up analysis, additional pages can be audited using the same methodology.", 10, False, GRAY)

    # Apply table formatting
    for table in doc.tables:
        table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(2)

    doc.save(output_path)
    print(f"   ✅ Report saved: {output_path}")


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="SEO & SEM Audit Tool — Crawl a URL and generate a .docx report using Claude or Gemini",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""
        Examples:
          # Using Claude (default)
          python seo_audit.py https://example.com --api-key sk-ant-api03-...

          # Using Gemini
          python seo_audit.py https://example.com --provider gemini --api-key AIzaSy...

          # Auto-detect provider from key format
          python seo_audit.py https://example.com -k AIzaSy...

          # Use environment variables
          export ANTHROPIC_API_KEY="sk-ant-..."   # for Claude
          export GEMINI_API_KEY="AIzaSy..."        # for Gemini
          python seo_audit.py https://example.com --provider gemini

          # Crawl only (no AI, free)
          python seo_audit.py https://example.com --crawl-only

        Requirements:
          pip install -r requirements.txt
        """)
    )
    parser.add_argument("url", help="URL to audit")
    parser.add_argument("--output", "-o", help="Output .docx file path (default: auto-generated)")
    parser.add_argument("--api-key", "-k", help="API key (Claude or Gemini)")
    parser.add_argument("--provider", "-p", choices=["claude", "gemini", "auto"],
                        default="auto", help="AI provider: claude, gemini, or auto-detect (default: auto)")
    parser.add_argument("--crawl-only", action="store_true", help="Only crawl, output JSON (skip AI analysis)")

    args = parser.parse_args()

    # Normalize URL
    url = args.url
    if not url.startswith("http"):
        url = "https://" + url

    # Generate output filename
    if args.output:
        output_path = args.output
    else:
        domain = urlparse(url).netloc.replace("www.", "").replace(".", "-")
        output_path = f"SEO-Report-{domain}-{datetime.now().strftime('%Y%m%d')}.docx"

    print("╔══════════════════════════════════════════════╗")
    print("║         SEO & SEM AUDIT TOOL                ║")
    print("║    AI-Powered Website Analysis               ║")
    print("║    Supports: Claude & Gemini                 ║")
    print("╚══════════════════════════════════════════════╝")

    # Step 1: Crawl
    crawler = SEOCrawler(url)
    crawl_data = crawler.crawl()

    if args.crawl_only:
        json_path = output_path.replace(".docx", ".json")
        with open(json_path, "w") as f:
            json.dump(crawl_data, f, indent=2, default=str)
        print(f"\n📁 Crawl data saved: {json_path}")
        return

    # Step 2: Resolve API key and provider
    api_key = args.api_key
    provider = args.provider

    # Fall back to environment variables (loaded from .env by dotenv at startup)
    if not api_key:
        if provider == "gemini":
            api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        elif provider == "claude":
            api_key = os.environ.get("ANTHROPIC_API_KEY")
        else:  # auto
            api_key = (os.environ.get("GEMINI_API_KEY")
                       or os.environ.get("GOOGLE_API_KEY")
                       or os.environ.get("ANTHROPIC_API_KEY"))

    if not api_key:
        print("\n❌ Error: API key required.")
        print("   Option 1 — Claude:  --api-key sk-ant-...  or  export ANTHROPIC_API_KEY=...")
        print("   Option 2 — Gemini:  --api-key AIzaSy...   or  export GEMINI_API_KEY=...")
        print("")
        print("   Get Claude key: https://console.anthropic.com/")
        print("   Get Gemini key: https://aistudio.google.com/apikey")
        sys.exit(1)

    # Auto-detect provider if needed
    if provider == "auto":
        provider = detect_provider(api_key)
        print(f"\n🔍 Auto-detected provider: {provider.upper()}")

    # Step 3: Load previous run for delta comparison
    history = load_history(url)
    previous_run = history[-1] if history else None
    if previous_run:
        print(f"\n📈 Previous run found: {previous_run['date']} — Score: {previous_run['overall_score']}/100")

    # Step 4: AI Analysis
    if provider == "gemini":
        analysis = analyze_with_gemini(crawl_data, api_key)
    else:
        analysis = analyze_with_claude(crawl_data, api_key)

    # Step 5: Compute deterministic scores from fixed rubric
    print("\n📐 Computing scores from fixed rubric...")
    analysis = compute_scores(analysis)

    # Step 6: Save to history
    save_history(url, analysis)

    # Step 7: Generate Report
    generate_report(analysis, url, output_path, previous_run=previous_run)

    print(f"\n🎉 Done! Report: {output_path}")
    print(f"   Provider: {provider.upper()}")
    print(f"   Overall Score: {analysis.get('overall_score', 'N/A')}/100 ({analysis.get('overall_grade', 'N/A')})")
    if previous_run:
        diff = analysis.get("overall_score", 0) - previous_run.get("overall_score", 0)
        sign = "+" if diff >= 0 else ""
        print(f"   vs Previous:   {previous_run['overall_score']}/100  ({sign}{diff} points)")


if __name__ == "__main__":
    main()